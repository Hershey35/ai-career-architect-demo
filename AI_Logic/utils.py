from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import io

def export_docx(resume_text, cover_letter_text):
    doc = Document()
    doc.add_heading('Tailored Resume', level=1)
    doc.add_paragraph(resume_text)

    doc.add_page_break()
    doc.add_heading('Cover Letter', level=1)
    doc.add_paragraph(cover_letter_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_pdf(resume_text, cover_letter_text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("<b>Tailored Resume</b>", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph(resume_text.replace("\n", "<br/>"), styles['Normal']))
    story.append(Spacer(1, 24))

    story.append(Paragraph("<b>Cover Letter</b>", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph(cover_letter_text.replace("\n", "<br/>"), styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer
